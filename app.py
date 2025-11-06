import os
import io
import json
import textwrap
import streamlit as st
from bs4 import BeautifulSoup
from readabilipy import simple_json_from_html_string
from docx import Document
import requests

# ---- Import your existing modules ----
from modules.fetch import fetch_article
from modules.nlp import extract_keywords_entities
from modules.serp import top_competitors, forum_questions
from modules.internal_links import suggest_internal_links, build_site_index
from modules.ai_writer import generate_draft

st.set_page_config(page_title="🧠 Mentalyc SEO Optimizer", layout="wide")

# ---------------- STYLE ----------------
BOLD_HEADER_STYLE = "padding:12px;border-radius:8px;color:white;font-weight:600;margin-top:24px;margin-bottom:6px;"

COLOR = {
    "keywords": "background:#7b47ff;",
    "competitors": "background:#ffb300;",
    "faqs": "background:#007bff;",
    "internal": "background:#28a745;",
    "draft": "background:#222;",
}

# ---------------- SIDEBAR ----------------
with st.sidebar:
    st.header("Settings")
    openai_key = st.sidebar.text_input("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", type="password")
serper_key = st.sidebar.text_input("5cb974df5481dfb855ca9c1878a7a1f71c2bd3836f6e72761316fba01bac7c8c", type="password")
site_base = st.sidebar.text_input("Site Base URL", value="https://www.mentalyc.com")
    model = st.selectbox("OpenAI Model", ["gpt-4.1", "gpt-4.1-mini"], index=1)
    gl = st.selectbox("Search Country (gl)", ["us", "in", "au", "gb", "ca"], index=0)
    hl = st.selectbox("Search Lang (hl)", ["en"], index=0)
    fast_mode = st.toggle("⚡ Speed Mode (skip heavy models & reduce API calls)", value=True, help="Skips embeddings, limits SERP calls, truncates text context for fastest startup.")

st.title("🧠 Mentalyc SEO Optimization Assistant — Bold UI + DOCX Export")("🧠 Mentalyc SEO Optimization Assistant — Bold UI + DOCX Export")
st.write("Enter Mentalyc blog URLs below to auto-generate SEO improvements.")

example_urls = "\n".join([
    "https://www.mentalyc.com/blog/psychological-assessment-report",
    "https://www.mentalyc.com/blog/icd-10-code-for-depression",
]))
urls_text = st.text_area("Target URLs (one per line)", height=140, value=example_urls)

if st.button("▶️ Run Optimization", type="primary"):
    urls = [u.strip() for u in urls_text.splitlines() if u.strip()]

    if not openai_key or not serper_key:
        st.error("Please provide both OpenAI and Serper API keys.")
        st.stop()

    blog_index = build_site_index(site_base) if not fast_mode else []
    progress = st.progress(0.0)

    for i, url in enumerate(urls, start=1):
        st.markdown("---")
        st.subheader(f"🔍 {url}")

        # Fetch + Analyze
        src = fetch_article(url)
        # In fast mode, skip spaCy NER/keybert if they slow down
kw = extract_keywords_entities(src.get("text", "")) if not fast_mode else {
    "primary_keyword": (src.get("title") or "").split("|")[0][:60],
    "secondary_keywords": [],
    "entities": []
}
        primary_q = kw.get("primary_keyword") or src.get("title", "")
        # Limit SERP calls in fast mode
comps = top_competitors(primary_q, serper_key, exclude_domain="mentalyc.com", gl=gl, hl=hl) if not fast_mode else []
        faqs = forum_questions(primary_q, serper_key, extras=kw.get("secondary_keywords", []), gl=gl, hl=hl) if not fast_mode else [], gl=gl, hl=hl)
        # Fast internal links: simple slug similarity when fast_mode
if fast_mode:
    from difflib import SequenceMatcher
    internals = []
else:
    internals = suggest_internal_links(url, src.get("text", ""), blog_index, n=6), blog_index, n=6)

        pack = {
            "source": src,
            "keywords": kw,
            "competitors": comps,
            "faqs_seed": faqs,
            "internal_suggestions": internals,
        }

        # Truncate source text in fast mode to keep prompt small
if fast_mode and pack["source"].get("text"):
    pack["source"]["text"] = pack["source"]["text"][:4000]

draft = generate_draft(pack, openai_key, model=model)
        pack["draft"] = draft

        # ---- Output UI ----
        st.markdown(f"<div style='{BOLD_HEADER_STYLE}{COLOR['keywords']}'>🟣 KEYWORDS & ENTITIES</div>", unsafe_allow_html=True)
        st.write(f"**Primary:** {kw.get('primary_keyword','')}\n\n**Secondary:** {', '.join(kw.get('secondary_keywords',[]))}")
        st.write(f"**Entities:** {', '.join(kw.get('entities',[]))}")

        st.markdown(f"<div style='{BOLD_HEADER_STYLE}{COLOR['competitors']}'>🟡 COMPETITORS</div>", unsafe_allow_html=True)
        for c in comps:
            st.write(f"- [{c['title']}]({c['url']})")

        st.markdown(f"<div style='{BOLD_HEADER_STYLE}{COLOR['faqs']}'>🔵 REAL USER FAQs</div>", unsafe_allow_html=True)
        for q in faqs:
            st.write(f"- {q}")

        st.markdown(f"<div style='{BOLD_HEADER_STYLE}{COLOR['internal']}'>🟢 INTERNAL LINK SUGGESTIONS</div>", unsafe_allow_html=True)
        for ln in internals:
            st.write(f"- {ln['target_url']} (score {ln['score']:.2f})")

        st.markdown(f"<div style='{BOLD_HEADER_STYLE}{COLOR['draft']}'>✍️ SEO DRAFT</div>", unsafe_allow_html=True)

        if isinstance(draft, dict):
            st.write(f"**Meta Title:** {draft.get('meta_title','')}")
            st.write(f"**Meta Description:** {draft.get('meta_description','')}")

            if draft.get("h2s"):
                st.write("### H2 Outline")
                st.write("\n".join([f"- {h}" for h in draft["h2s"]]))

            st.write("---")
            st.markdown(draft.get("body",""))

            # ---- DOCX Export (D1: direct download) ----
            def create_docx(d):
                doc = Document()
                doc.add_heading(d.get("meta_title",""), level=1)
                doc.add_paragraph(d.get("meta_description",""))
                for h in d.get("h2s", []):
                    doc.add_heading(h, level=2)
                body = d.get("body","" ).replace("\n","\n\n")
                doc.add_paragraph(body)
                f = io.BytesIO()
                doc.save(f)
                return f.getvalue()

            docx_data = create_docx(draft)
            st.download_button("⬇️ Download DOCX", data=docx_data, file_name=f"{url.split('/')[-1]}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        else:
            st.code(draft)

        progress.progress(i / len(urls))

    st.success("✅ Finished!")
